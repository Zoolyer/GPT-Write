import time

from docx import Document
from docx.shared import Pt, Inches
import re
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
import docx
from sel import seg
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def process_txt_modified(input_filename, output_filename):

    # 读取文本文件内容
    with open(input_filename, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    # 定义一个函数来检查行中是否有括号内容
    def has_brackets_content(line, tag):
        pattern = re.compile(r'<'+ tag + r'>.*（.*?）.*</' + tag + r'>')
        return pattern.search(line) is not None

        # 定义一个函数来检查行中是否有普通括号内容
    def has_regular_brackets_content(line):
        pattern = re.compile(r'.*\(.*\).*')
        return pattern.search(line) is not None

        # 定义一个函数来检查行中是否有特定的括号内容
    def has_specific_brackets_content(line, start_bracket, end_bracket):
        pattern = re.compile(r'.*' + re.escape(start_bracket) + r'(继续|续)' + re.escape(end_bracket) + r'.*')
        return pattern.search(line) is not None

    def remove_tags_from_line(line):
        """移除指定的标签并返回移除的标签和没有标签的内容"""
        tags_to_remove = ["<t1>", "</t1>", "<t2>", "</t2>", "<t3>", "</t3>"]
        saved_tags = [tag for tag in tags_to_remove if tag in line]
        for tag in tags_to_remove:
            line = line.replace(tag, "")
        return saved_tags, line.strip()

    def remove_starting_index(line):
        split_content = line.split(" ")
        if len(split_content) > 1:
            return split_content[-1].strip()
        return line

    # 删除包含（内容）的<t1>, <t2>, <t3>标签行和包含普通括号的行
    cleaned_lines = []
    for line in lines:
        saved_tags, line_without_tags = remove_tags_from_line(line)

        if any(tag in saved_tags for tag in ["<t1>", "<t2>", "<t3>"]):
            # 如果存在特定的括号内容，则跳过此行
            if any(has_specific_brackets_content(line_without_tags, bracket[0], bracket[1]) for bracket in
                   [("（", "）"), ("(", ")")]):
                continue
            cleaned_content = remove_starting_index(line_without_tags)
            # 将保存的标签添加回去
            cleaned_line = "".join(saved_tags[0]) + cleaned_content +saved_tags[1]
            # print(cleaned_line)
            cleaned_lines.append(cleaned_line)
        else:
            if any(has_specific_brackets_content(line_without_tags, bracket[0], bracket[1]) for bracket in
                   [("（", "）"), ("(", ")")]):
                continue
            # print("line:"+line)
            cleaned_lines.append(line)

    # 合并处理
    def merge_tags(content, tag):
    # 正则表达式模式匹配两个相邻的、内容相同的同一级别的标签
        pattern = re.compile(r'<' + tag + r'>(.*?)</' + tag + r'>\s*<' + tag + r'>\1</' + tag + r'>')
        while pattern.search(content):
            # 当匹配到两个相同内容的标签时，删除第二个标签
            content = pattern.sub(r'<'+ tag + r'>\1', content)
        return content


    text_merged = ''.join(cleaned_lines)
    for tag in ["t1", "t2", "t3"]:
        text_merged = merge_tags(text_merged, tag)

    # 写入到输出文件
    with open(output_filename, 'w', encoding='utf-8') as file:
        file.write(text_merged)

    return f"Processed content written to {output_filename}"
def remove_num_format(text):
    pattern = r'\d+ / \d'
    return re.sub(pattern, '', text)
def process_content(content, with_numbering=False):
    content = remove_num_format(content)  # Remove (num / num) format
    elements = []

    # 标题计数器
    count_t1 = 0
    count_t2 = 0
    count_t3 = 0

    while True:
        match = re.search(r'<(.*?)>(.*?)</\1>', content, re.DOTALL)
        if not match:
            # 如果没有找到标签，则处理剩余的文本
            for line in content.strip().split('\n'):
                if line.strip():
                    elements.append(('p', line.strip()))
            break

        # 添加到标签之前的文本（如果有的话）
        start_tag, end_tag = match.span()
        if start_tag > 0:
            for line in content[:start_tag].strip().split('\n'):
                if line.strip():
                    elements.append(('p', line.strip()))

        # 处理标签内的文本
        tag, text = match.groups()
        text_lines = text.strip().split('\n')
        for line in text_lines:
            if tag == "t1":
                if with_numbering:
                    count_t1 += 1
                    count_t2 = 0
                    count_t3 = 0
                    elements.append(('t1', f"{count_t1}. {line.strip()}"))
                else:
                    elements.append(('t1', line.strip()))
            elif tag == "t2":
                if with_numbering:
                    count_t2 += 1
                    count_t3 = 0
                    elements.append(('t2', f"{count_t1}.{count_t2}. {line.strip()}"))
                else:
                    elements.append(('t2', line.strip()))
            elif tag == "t3":
                if with_numbering:
                    count_t3 += 1
                    elements.append(('t3', f"{count_t1}.{count_t2}.{count_t3}. {line.strip()}"))
                else:
                    elements.append(('t3', line.strip()))
            elif tag == "b":
                elements.append(('b', line.strip()))
            else:
                elements.append(('p', line.strip()))

        # 从原始文本中删除处理过的部分
        content = content[end_tag:]

    return elements


def set_paragraph_style(p, font_name, font_size):
    if p.runs:
        run = p.runs[0]
        run.font.name = font_name
        run.font.size = Pt(font_size)

def set_run_style(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def txt_to_word(abstract_input, text_input, conclusion_input, output_docx,title):
    doc = Document()


    # 添加两个空段落
    for i in range(6):
        doc.add_paragraph()

    # 添加标题段落并设置其样式
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    set_run_style(title_run, '黑体', 28)
    title_run.bold = True
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐



    doc.add_section()
    # 获取 Heading 1 样式
    heading1 = doc.styles['Heading 1']
    # 修改 Heading 1 样式的加粗属性
    heading1.font.bold = False
    # 获取 Heading 1 样式
    heading2 = doc.styles['Heading 2']
    # 修改 Heading 1 样式的加粗属性
    heading2.font.bold = False
    # 获取 Heading 1 样式
    heading3 = doc.styles['Heading 3']
    # 修改 Heading 1 样式的加粗属性
    heading3.font.bold = False
    # Process abstract_input
    with open(abstract_input, 'r', encoding='utf-8') as file:
        content = file.read()
    elements = process_content(content)
    for tag, text in elements:
        if not text.strip():  # Skip empty lines
            continue

        if "关键词：" in text or "关键词:" in text:
            index_colon = text.find("关键词：") if "关键词：" in text else text.find("关键词:")
            before_keyword = text[:index_colon]
            after_keyword = text[index_colon:]

            if before_keyword.strip():
                p = doc.add_paragraph()
                run = p.add_run(before_keyword.strip())
                run.font.color.rgb = RGBColor(0, 0, 0)  # RGB for black
                set_run_style(run, '宋体', 10.5)
                p.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进

            p = doc.add_paragraph()
            p.style = doc.styles['Normal']  # 基于“正文”样式
            run = p.add_run("关键词")
            run.bold = True
            set_run_style(run, '宋体', 10.5)
            run = p.add_run(after_keyword[3:])
            set_run_style(run, '宋体', 10.5)
            p.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进

        else:
            if "引言" in text and tag == "t1":
                doc.add_page_break()

                # 添加目录标题到第二部分的第二页并居中
                paragraph = doc.add_paragraph()
                paragraph.style = doc.styles['Heading 1']
                run = paragraph.add_run("目录")
                run.font.color.rgb = RGBColor(0, 0, 0)  # RGB for black
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
                set_run_style(run, '黑体', 14)


                doc.add_page_break()
                doc.add_section()

            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 0)  # RGB for black
            if tag == "t1":
                p.style = doc.styles['Heading 1']
                set_run_style(run, '黑体', 12)
                p.paragraph_format.space_before = Pt(18)  # 段前0.5行
                p.paragraph_format.space_after = Pt(18)  # 段后0.5行
                if "摘要" in text or "引言" in text:
                    p.paragraph_format.alignment = 1  # Set to center
            else:
                p.style = doc.styles['Normal']  # 基于“正文”样式
                set_run_style(run, '宋体', 10.5)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing = Pt(18)  # 1.5倍行距
                p.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进
    process_txt_modified(text_input,"text_input2.txt")
    # Process text_input with numbering
    with open("text_input2.txt", 'r', encoding='utf-8') as file:
        content = file.read()
    elements = process_content(content, with_numbering=True)


    for tag, text in elements:
        if not text.strip():  # Skip empty lines
            continue
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 0)  # RGB for black
        if tag == "t1":
            p.style = doc.styles['Heading 1']
            set_run_style(run, '黑体', 12)
            p.paragraph_format.space_before = Pt(18)  # 段前0.5行
            p.paragraph_format.space_after = Pt(18)  # 段后0.5行
            p.paragraph_format.line_spacing = Pt(20)  # 1.5倍行距
        elif tag == "t2":
            p.style = doc.styles['Heading 2']
            set_run_style(run, '黑体', 10.5)
            p.paragraph_format.line_spacing = Pt(20)  # 1.5倍行距
            p.paragraph_format.space_after = Pt(10)  # 段后0.5行
        elif tag == "t3":
            p.style = doc.styles['Heading 3']
            set_run_style(run, '宋体', 10.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

            p.paragraph_format.line_spacing = Pt(20)  # 1.5倍行距
            p.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进
        else:
            p.style = doc.styles['Normal']  # 基于“正文”样式
            set_run_style(run, '宋体', 10.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = Pt(18)  # 1.5倍行距
            p.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进




    # Process conclusion_input
    with open(conclusion_input, 'r', encoding='utf-8') as file:
        content = file.read()
    elements = process_content(content)
    for tag, text in elements:
        if "总结" in text and tag == "t1":
            doc.add_page_break()
        if not text.strip():  # Skip empty lines
            continue
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 0)  # RGB for black
        if tag == "t1":
            p.style = doc.styles['Heading 1']
            set_run_style(run, '黑体', 12)
            if "总结" in text:
                p.paragraph_format.alignment = 1  # Set to center
            p.paragraph_format.space_before = Pt(18)  # 段前0.5行
            p.paragraph_format.space_after = Pt(18)  # 段后0.5行
        else:
            p.style = doc.styles['Normal']  # 基于“正文”样式
            set_run_style(run, '宋体', 10.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = Pt(18)  # 1.5倍行距
            p.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = Pt(18)  # 1.5倍行距

    # 断开第二节与第一节的页脚链接
    doc.sections[0].footer.is_linked_to_previous = False
    doc.sections[1].footer.is_linked_to_previous = False
    doc.sections[2].footer.is_linked_to_previous = False
    # 定位到第一节的页脚
    footer = doc.sections[0].footer

    # 清空第一节的页脚内容
    for paragraph in footer.paragraphs:
        paragraph.clear()


    # 定位到第二节的页脚

    section = doc.sections[1]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.starting_page_number = 1


    footer = section.footer

    # 找到或创建页码需要插入的段落
    if len(footer.paragraphs) == 0:
        paragraph = footer.add_paragraph()
    else:
        paragraph = footer.paragraphs[0]

    # 在段落中插入页码
    add_page_number(paragraph)

    section = doc.sections[1]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.starting_page_number = 1


    doc.save(output_docx)

def add_page_number(paragraph):
    """向段落中添加页码。"""
    run = paragraph.add_run()
    fldChar1 = docx.oxml.OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = docx.oxml.OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = docx.oxml.OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
def add_footer_with_page_number(section, format_type, start_page_number):
    footer = section.footer
    if not footer.paragraphs:
        paragraph = footer.add_paragraph()
    else:
        paragraph = footer.paragraphs[0]
        # Clear any content in the footer paragraph to avoid duplicating page numbers
        for element in paragraph._element:
            paragraph._element.remove(element)

    field_code = 'PAGE'
    if format_type == "roman":
        field_code += ' \* ROMAN '

    # Add field for page number
    fldSimple = docx.oxml.OxmlElement('w:fldSimple')
    fldSimple.set(docx.oxml.ns.qn('w:instr'), field_code)

    run = docx.oxml.OxmlElement('w:r')
    fldSimple.append(run)

    # Set the color of the page number to black
    color = docx.oxml.OxmlElement('w:color')
    color.set(docx.oxml.ns.qn('w:val'), "000000")
    run.append(color)

    t = docx.oxml.OxmlElement('w:t')
    run.append(t)
    t.text = "1"  # The actual page number will replace this 1

    paragraph._element.append(fldSimple)

    # Set the alignment of the paragraph to center
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set section start page number
    if start_page_number:
        if format_type == "roman":
            section.start_page_number = int_to_roman(start_page_number)
        else:
            section.start_page_number = start_page_number





def int_to_roman(num):
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
        ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
        ]
    roman_num = ''
    i = 0
    while  num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num
# 调用函数
if __name__ == "__main__":
    txt_to_word('abstract_input.txt', 'text_input.txt', 'conclusion_input.txt', 'output.docx',"我是标题")
    time.sleep(1)
    # seg()